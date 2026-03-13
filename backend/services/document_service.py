"""
Document service: file storage, text extraction, and lifecycle management.

Handles uploading documents to Supabase Storage, extracting text content
from PDF, DOCX, and XLSX files, and managing document metadata in the
documents table.
"""
from __future__ import annotations

import io
import logging
import time
import uuid
from datetime import date
from typing import Any

from services.supabase_service import get_supabase
from services.memory_service import _ensure_bucket, _storage, append_daily_log, ensure_daily_file
from services.embedding_service import index_memory_file
from services.openai_service import summarize_document
from services import google_workspace_service

log = logging.getLogger(__name__)


# ── Allowed file types ────────────────────────────────────────────────

ALLOWED_EXTENSIONS = {"pdf", "docx", "xlsx", "doc", "xls"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB
GOOGLE_DRIVE_SOURCE = "google_drive"


def _document_storage_path(user_id: str, filename: str) -> str:
    """Build a unique Storage path for a document."""
    unique = uuid.uuid4().hex[:8]
    safe_name = filename.replace(" ", "_")
    return f"{user_id}/documents/{unique}_{safe_name}"


def get_file_extension(filename: str) -> str:
    """Extract and normalise the file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext


def _normalise_drive_version(value: Any) -> str:
    return str(value or "")


def _document_source_update_payload(source: dict[str, Any] | None) -> dict[str, Any]:
    if not source:
        return {}
    return {
        "source": source.get("source", "upload"),
        "source_external_id": source.get("external_id"),
        "source_web_url": source.get("web_url"),
        "source_version": _normalise_drive_version(source.get("version")),
        "source_modified_time": source.get("modified_time"),
        "source_checksum": source.get("checksum"),
    }


# ── Text extraction ───────────────────────────────────────────────────

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(file_bytes))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text.strip())
    return "\n\n".join(parts)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file (Word / Google Docs export)."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_bytes))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def extract_text_from_xlsx(file_bytes: bytes) -> str:
    """Extract text from an XLSX/XLS file."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts: list[str] = []

    for sheet in wb.sheetnames:
        ws = wb[sheet]
        parts.append(f"## Sheet: {sheet}")

        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() if c is not None else "" for c in row]
            non_empty = [c for c in cells if c]
            if non_empty:
                parts.append(" | ".join(cells))

    wb.close()
    return "\n\n".join(parts)


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Route to the correct extractor based on file type."""
    if file_type in ("pdf",):
        return extract_text_from_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file_bytes)
    elif file_type in ("xlsx", "xls"):
        return extract_text_from_xlsx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ── Storage operations ────────────────────────────────────────────────

def store_document(user_id: str, filename: str, file_bytes: bytes, content_type: str) -> str:
    """
    Upload a document to Supabase Storage.

    Returns the storage path.
    """
    _ensure_bucket()
    storage = _storage()
    path = _document_storage_path(user_id, filename)

    storage.upload(
        path,
        file_bytes,
        {"content-type": content_type},
    )
    return path


def ingest_document_upload(
    user_id: str,
    filename: str,
    file_bytes: bytes,
    content_type: str = "application/octet-stream",
    source: dict[str, Any] | None = None,
) -> dict:
    """
    Persist and process an uploaded document payload.

    Mirrors the normal /documents/upload flow so callers (e.g. Inbox
    attachment imports) trigger the same extraction, indexing, and daily
    memory summary behavior.
    """
    ext = get_file_extension(filename)
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: .{ext}. Accepted: {', '.join(sorted(ALLOWED_EXTENSIONS))}"
        )

    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(
            f"File too large. Maximum size is {MAX_FILE_SIZE // (1024 * 1024)}MB."
        )

    sb = get_supabase()
    started = time.monotonic()
    log.info(
        "document_ingest_start user=%s filename=%s size=%d ext=%s",
        user_id,
        filename,
        len(file_bytes),
        ext,
    )

    storage_path = store_document(user_id, filename, file_bytes, content_type)
    insert_payload = {
        "user_id": user_id,
        "filename": filename,
        "file_type": ext,
        "file_size": len(file_bytes),
        "storage_path": storage_path,
        "status": "processing",
    }
    insert_payload.update(_document_source_update_payload(source))
    result = sb.table("documents").insert(insert_payload).execute()
    doc = result.data[0]

    try:
        process_document(user_id, doc["id"], storage_path, ext)
        refreshed = sb.table("documents").select("*").eq("id", doc["id"]).execute()
        doc = refreshed.data[0] if refreshed.data else doc
    except Exception:
        log.exception("process_document failed for doc=%s user=%s", doc["id"], user_id)
        refreshed = sb.table("documents").select("*").eq("id", doc["id"]).execute()
        doc = refreshed.data[0] if refreshed.data else doc

    elapsed = (time.monotonic() - started) * 1000
    log.info(
        "document_ingest_done user=%s doc=%s status=%s duration_ms=%.0f",
        user_id,
        doc.get("id"),
        doc.get("status"),
        elapsed,
    )
    return doc


def ingest_google_drive_document(
    user_id: str,
    credentials_json: str,
    file_id: str,
) -> tuple[dict, str | None]:
    """Ingest a Google Drive file by downloading/exporting to supported binary."""
    payload, refreshed = google_workspace_service.drive_download_for_ingestion(
        credentials_json,
        file_id=file_id,
    )
    ingestion = payload["ingestion"]
    file_meta = payload["file"]
    doc = ingest_document_upload(
        user_id=user_id,
        filename=ingestion["filename"],
        file_bytes=ingestion["bytes"],
        content_type=ingestion["content_type"],
        source={
            "source": GOOGLE_DRIVE_SOURCE,
            "external_id": file_meta.get("id"),
            "web_url": file_meta.get("webViewLink"),
            "version": file_meta.get("version"),
            "modified_time": file_meta.get("modifiedTime"),
            "checksum": file_meta.get("md5Checksum"),
        },
    )
    return doc, refreshed


def download_document(storage_path: str) -> bytes:
    """Download a document from Supabase Storage."""
    _ensure_bucket()
    storage = _storage()
    data = storage.download(storage_path)
    return data if isinstance(data, bytes) else data.encode("utf-8")


def delete_document_file(storage_path: str) -> None:
    """Delete a document from Supabase Storage."""
    _ensure_bucket()
    storage = _storage()
    try:
        storage.remove([storage_path])
    except Exception:
        pass  # best-effort deletion


# ── Full lifecycle ────────────────────────────────────────────────────

def process_document(user_id: str, document_id: str, storage_path: str, file_type: str) -> None:
    """
    Download, extract text, index into memory_chunks, and update status.

    This is called after the initial upload. It:
    1. Downloads the file from Storage
    2. Extracts text content
    3. Stores extracted text in the documents table
    4. Indexes content into memory_chunks for RAG
    5. Updates status to 'ready' (or 'error')
    """
    sb = get_supabase()

    try:
        # 1. Download
        file_bytes = download_document(storage_path)

        # 2. Extract text
        text = extract_text(file_bytes, file_type)

        if not text.strip():
            log.warning(
                "document_extract_empty user=%s doc=%s type=%s path=%s",
                user_id,
                document_id,
                file_type,
                storage_path,
            )
            sb.table("documents").update({
                "status": "error",
                "extracted_text": "",
            }).eq("id", document_id).execute()
            return

        # 3. Save extracted text to documents table
        sb.table("documents").update({
            "extracted_text": text,
            "status": "ready",
        }).eq("id", document_id).execute()

        # 4. Index into memory_chunks for RAG search
        # Use source_file format: "documents/<filename>"
        result = sb.table("documents").select("filename").eq("id", document_id).execute()
        if result.data:
            source_file = f"documents/{result.data[0]['filename']}"
            index_memory_file(user_id, source_file, text)

        # 5. Generate summary and save to daily memory log
        #    Wrapped in its own try/except so summarisation failure never
        #    marks the document as 'error' — the document is already ready.
        try:
            if result.data:
                filename = result.data[0]["filename"]
                summary = summarize_document(text, filename)
                if summary:
                    memory_entry = (
                        f"## Document uploaded: {filename}\n\n"
                        f"{summary}"
                    )
                    ensure_daily_file(user_id)
                    updated_daily = append_daily_log(user_id, memory_entry)

                    # Re-index the daily log so the summary is searchable
                    today_source = f"daily/{date.today().isoformat()}.md"
                    index_memory_file(user_id, today_source, updated_daily)
        except Exception:
            log.exception(
                "document_summary_failed_non_blocking user=%s doc=%s",
                user_id,
                document_id,
            )

    except Exception as e:
        # Mark as error
        log.exception(
            "process_document_failed user=%s doc=%s type=%s path=%s",
            user_id,
            document_id,
            file_type,
            storage_path,
        )
        sb.table("documents").update({
            "status": "error",
        }).eq("id", document_id).execute()
        raise e


def refresh_google_drive_document_if_stale(
    user_id: str,
    document: dict[str, Any],
    credentials_json: str,
) -> tuple[dict[str, Any], str | None]:
    """
    Lazy revalidate linked Drive docs and refresh local binary if remote changed.
    """
    if document.get("source") != GOOGLE_DRIVE_SOURCE:
        return document, None
    file_id = (document.get("source_external_id") or "").strip()
    if not file_id:
        return document, None

    remote, token_refresh = google_workspace_service.drive_get_file_metadata(credentials_json, file_id)
    remote_version = _normalise_drive_version(remote.get("version"))
    remote_modified = remote.get("modifiedTime")
    remote_checksum = remote.get("md5Checksum")

    stale = any([
        remote_version and remote_version != _normalise_drive_version(document.get("source_version")),
        remote_modified and remote_modified != document.get("source_modified_time"),
        remote_checksum and remote_checksum != document.get("source_checksum"),
    ])
    if not stale:
        return document, token_refresh

    payload, download_refresh = google_workspace_service.drive_download_for_ingestion(
        credentials_json,
        file_id=file_id,
    )
    ingestion = payload["ingestion"]
    file_meta = payload["file"]
    new_filename = ingestion["filename"]
    new_type = get_file_extension(new_filename)
    if new_type not in ALLOWED_EXTENSIONS:
        raise ValueError(f"Unsupported linked Drive file type: .{new_type}")
    if len(ingestion["bytes"]) > MAX_FILE_SIZE:
        raise ValueError(
            f"Linked Drive file too large. Maximum size is {MAX_FILE_SIZE // (1024 * 1024)}MB."
        )

    sb = get_supabase()
    old_storage_path = document.get("storage_path")
    old_source_file = f"documents/{document.get('filename', '')}"
    new_storage_path = store_document(
        user_id,
        new_filename,
        ingestion["bytes"],
        ingestion["content_type"],
    )

    sb.table("documents").update({
        "filename": new_filename,
        "file_type": new_type,
        "file_size": len(ingestion["bytes"]),
        "storage_path": new_storage_path,
        "status": "processing",
        "source_web_url": file_meta.get("webViewLink"),
        "source_version": _normalise_drive_version(file_meta.get("version")),
        "source_modified_time": file_meta.get("modifiedTime"),
        "source_checksum": file_meta.get("md5Checksum"),
    }).eq("id", document["id"]).eq("user_id", user_id).execute()

    process_document(user_id, document["id"], new_storage_path, new_type)

    if old_storage_path:
        delete_document_file(old_storage_path)
    if document.get("filename") and document.get("filename") != new_filename:
        sb.table("memory_chunks").delete().eq("user_id", user_id).eq("source_file", old_source_file).execute()

    refreshed_doc_result = (
        sb.table("documents")
        .select("*")
        .eq("id", document["id"])
        .eq("user_id", user_id)
        .limit(1)
        .execute()
    )
    final_doc = refreshed_doc_result.data[0] if refreshed_doc_result.data else document
    return final_doc, (download_refresh or token_refresh)


def delete_document_full(user_id: str, document_id: str) -> None:
    """
    Delete a document completely: Storage file, memory_chunks, and DB row.
    """
    sb = get_supabase()

    # Get document info
    result = sb.table("documents").select("*").eq("id", document_id).eq("user_id", user_id).execute()
    if not result.data:
        raise ValueError("Document not found")

    doc = result.data[0]

    # 1. Delete from Storage
    delete_document_file(doc["storage_path"])

    # 2. Delete memory chunks
    source_file = f"documents/{doc['filename']}"
    sb.table("memory_chunks") \
        .delete() \
        .eq("user_id", user_id) \
        .eq("source_file", source_file) \
        .execute()

    # 3. Delete DB row
    sb.table("documents").delete().eq("id", document_id).execute()
